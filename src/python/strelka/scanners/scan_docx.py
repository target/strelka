import io
import zipfile

import docx
from bs4 import BeautifulSoup

from strelka import strelka


class ScanDocx(strelka.Scanner):
    """Collects metadata and extracts text from docx files.

    Options:
        extract_text: Boolean that determines if document text should be
            extracted as a child file.
            Defaults to False.
    """

    def scan(self, data, file, options, expire_at):
        extract_text = options.get("extract_text", False)
        with io.BytesIO(data) as docx_io:
            try:
                docx_doc = docx.Document(docx_io)
                self.event["author"] = docx_doc.core_properties.author
                self.event["category"] = docx_doc.core_properties.category
                self.event["comments"] = docx_doc.core_properties.comments
                self.event["content_status"] = docx_doc.core_properties.content_status
                if docx_doc.core_properties.created is not None:
                    self.event["created"] = docx_doc.core_properties.created.isoformat()
                self.event["identifier"] = docx_doc.core_properties.identifier
                self.event["keywords"] = docx_doc.core_properties.keywords
                self.event["language"] = docx_doc.core_properties.language
                self.event["last_modified_by"] = (
                    docx_doc.core_properties.last_modified_by
                )
                if docx_doc.core_properties.last_printed is not None:
                    self.event["last_printed"] = (
                        docx_doc.core_properties.last_printed.isoformat()
                    )
                if docx_doc.core_properties.modified is not None:
                    self.event["modified"] = (
                        docx_doc.core_properties.modified.isoformat()
                    )
                self.event["revision"] = docx_doc.core_properties.revision
                self.event["subject"] = docx_doc.core_properties.subject
                self.event["title"] = docx_doc.core_properties.title
                self.event["version"] = docx_doc.core_properties.version
                self.event["font_colors"] = [""]
                self.event["word_count"] = 0
                self.event["image_count"] = 0

                for paragraph in docx_doc.paragraphs:
                    soup = BeautifulSoup(paragraph.paragraph_format.element.xml, "xml")
                    color_list = soup.select("color")

                    for color_xml in color_list:
                        color = color_xml.attrs["w:val"]
                        if color not in self.event["font_colors"]:
                            self.event["font_colors"].append(color)

                    image_list = soup.select("pic")

                    for images in image_list:
                        if images.attrs["xmlns:pic"]:
                            self.event["image_count"] += 1

                    para_words = paragraph.text.split(" ")

                    if "" not in para_words:
                        self.event["word_count"] += len(para_words)

                if "FFFFFF" in self.event["font_colors"]:
                    self.event["white_text_in_doc"] = True

                if extract_text:
                    text = ""
                    for paragraph in docx_doc.paragraphs:
                        text += f"{paragraph.text}\n"

                    # Send extracted file back to Strelka
                    self.emit_file(text.encode("utf-8"), name="text")

            except ValueError:
                self.flags.append("value_error")
            except zipfile.BadZipFile:
                self.flags.append("bad_zip")
            except strelka.ScannerTimeout:
                raise
            except Exception:
                self.flags.append("bad_doc")
